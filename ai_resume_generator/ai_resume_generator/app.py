from flask import Flask, render_template, request, send_file, jsonify
import requests
import json
from datetime import datetime
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io

app = Flask(__name__)

def ensure_directories():
    """Ensure all required directories exist"""
    directories = [
        'generated_documents',
        'templates',
        'static'
    ]
    
    for directory in directories:
        dir_path = Path(directory)
        dir_path.mkdir(parents=True, exist_ok=True)
        print(f"✅ Directory ensured: {dir_path.absolute()}")

ensure_directories()

class LocalLLM:
    def __init__(self, model_name="llama2:7b"):
        self.model_name = model_name
        self.base_url = "http://localhost:11434/api/generate"
    
    def generate_text(self, prompt, max_tokens=1000):
        """Generate text using local Ollama model"""
        try:
            payload = {
                "model": self.model_name,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "num_predict": max_tokens,
                    "temperature": 0.7
                }
            }
            
            response = requests.post(self.base_url, json=payload, timeout=120)
            
            if response.status_code == 200:
                result = response.json()
                return result.get('response', 'No response generated')
            else:
                return f"HTTP {response.status_code}: {response.text}"
                
        except requests.exceptions.ConnectionError:
            return "Connection failed: Is Ollama running? Start it with 'ollama serve'"
        except Exception as e:
            return f"Unexpected error: {str(e)}"

llm = LocalLLM()

def add_hyperlink(paragraph, text, url, style_name=None):
    """Add a hyperlink to a paragraph"""
    try:
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Add color and underline formatting
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        
        paragraph._p.append(hyperlink)
        return hyperlink
    except Exception as e:
        print(f"Warning: Could not create hyperlink: {e}")
        paragraph.add_run(text)
        return None

class ResumeTemplates:
    @staticmethod
    def create_modern_template(doc, data, page_limit):
        """Modern Professional Template"""
        # Header with name and contact
        header_table = doc.add_table(rows=2, cols=2)
        header_table.autofit = True
        
        # Name cell
        name_cell = header_table.cell(0, 0)
        name_para = name_cell.paragraphs[0]
        name_run = name_para.add_run(data.get('name', ''))
        name_run.font.size = Pt(20)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
        
        # Title cell
        title_cell = header_table.cell(0, 1)
        title_para = title_cell.paragraphs[0]
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        title_run = title_para.add_run(data.get('job_title', ''))
        title_run.font.size = Pt(14)
        title_run.font.italic = True
        
        # Contact info
        contact_cell = header_table.cell(1, 0)
        contact_para = contact_cell.paragraphs[0]
        contact_info = []
        if data.get('email'):
            contact_info.append(data.get('email'))
        if data.get('phone'):
            contact_info.append(data.get('phone'))
        if data.get('location'):
            contact_info.append(data.get('location'))
        contact_para.add_run(' | '.join(contact_info))
        
        # Social links
        social_cell = header_table.cell(1, 1)
        social_para = social_cell.paragraphs[0]
        social_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        if data.get('linkedin'):
            add_hyperlink(social_para, "LinkedIn", data.get('linkedin'))
            if data.get('github') or data.get('portfolio'):
                social_para.add_run(" | ")
        
        if data.get('github'):
            add_hyperlink(social_para, "GitHub", data.get('github'))
            if data.get('portfolio'):
                social_para.add_run(" | ")
        
        if data.get('portfolio'):
            add_hyperlink(social_para, "Portfolio", data.get('portfolio'))
        
        # Remove table borders
        ResumeTemplates._remove_table_borders(header_table)
        
        # Horizontal line
        doc.add_paragraph().add_run("_" * 80)
        
        return doc
    
    @staticmethod
    def create_classic_template(doc, data, page_limit):
        """Classic Professional Template"""
        # Centered header
        header = doc.add_paragraph()
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Name
        name_run = header.add_run(data.get('name', ''))
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.name = 'Times New Roman'
        
        header.add_run('\n')
        
        # Contact info
        contact_info = []
        if data.get('email'):
            contact_info.append(data.get('email'))
        if data.get('phone'):
            contact_info.append(data.get('phone'))
        if data.get('location'):
            contact_info.append(data.get('location'))
        
        if contact_info:
            contact_run = header.add_run(' • '.join(contact_info))
            contact_run.font.size = Pt(11)
        
        # Social links on new line
        if data.get('linkedin') or data.get('github') or data.get('portfolio'):
            header.add_run('\n')
            
            if data.get('linkedin'):
                add_hyperlink(header, "LinkedIn Profile", data.get('linkedin'))
                if data.get('github') or data.get('portfolio'):
                    header.add_run(" • ")
            
            if data.get('github'):
                add_hyperlink(header, "GitHub Profile", data.get('github'))
                if data.get('portfolio'):
                    header.add_run(" • ")
            
            if data.get('portfolio'):
                add_hyperlink(header, "Portfolio", data.get('portfolio'))
        
        # Horizontal line
        doc.add_paragraph("═" * 60).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        return doc
    
    @staticmethod
    def create_creative_template(doc, data, page_limit):
        """Creative Modern Template"""
        # Create a table for layout
        main_table = doc.add_table(rows=1, cols=2)
        main_table.autofit = False
        main_table.columns[0].width = Inches(2.0)
        main_table.columns[1].width = Inches(4.5)
        
        # Left column (sidebar)
        left_cell = main_table.cell(0, 0)
        left_cell.paragraphs[0].clear()
        
        # Name in sidebar
        name_para = left_cell.add_paragraph()
        name_run = name_para.add_run(data.get('name', ''))
        name_run.font.size = Pt(16)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
        # Add background color to left cell
        shading_elm = parse_xml(r'<w:shd {} w:fill="2E86AB"/>'.format(nsdecls('w')))
        left_cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Contact info in sidebar
        contact_para = left_cell.add_paragraph()
        contact_run = contact_para.add_run("CONTACT\n")
        contact_run.font.bold = True
        contact_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
        if data.get('email'):
            email_para = left_cell.add_paragraph()
            email_run = email_para.add_run(data.get('email'))
            email_run.font.size = Pt(9)
            email_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
        if data.get('phone'):
            phone_para = left_cell.add_paragraph()
            phone_run = phone_para.add_run(data.get('phone'))
            phone_run.font.size = Pt(9)
            phone_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
        # Social links in sidebar
        if data.get('linkedin') or data.get('github') or data.get('portfolio'):
            social_header = left_cell.add_paragraph()
            social_header_run = social_header.add_run("\nLINKS\n")
            social_header_run.font.bold = True
            social_header_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
            if data.get('linkedin'):
                linkedin_para = left_cell.add_paragraph()
                linkedin_run = linkedin_para.add_run("LinkedIn")
                linkedin_run.font.size = Pt(9)
                linkedin_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                linkedin_run.font.underline = True
            
            if data.get('github'):
                github_para = left_cell.add_paragraph()
                github_run = github_para.add_run("GitHub")
                github_run.font.size = Pt(9)
                github_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                github_run.font.underline = True
            
            if data.get('portfolio'):
                portfolio_para = left_cell.add_paragraph()
                portfolio_run = portfolio_para.add_run("Portfolio")
                portfolio_run.font.size = Pt(9)
                portfolio_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                portfolio_run.font.underline = True
        
        # Right column (main content)
        right_cell = main_table.cell(0, 1)
        right_cell.paragraphs[0].clear()
        
        # Job title
        title_para = right_cell.add_paragraph()
        title_run = title_para.add_run(data.get('job_title', ''))
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
        
        ResumeTemplates._remove_table_borders(main_table)
        
        return doc
    
    @staticmethod
    def create_minimal_template(doc, data, page_limit):
        """Minimal Clean Template"""
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(data.get('name', ''))
        name_run.font.size = Pt(24)
        name_run.font.bold = True
        name_run.font.name = 'Calibri'
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Job title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(data.get('job_title', ''))
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        title_para.space_after = Pt(12)
        
        # Contact info in one line
        contact_para = doc.add_paragraph()
        contact_items = []
        
        if data.get('email'):
            contact_items.append(data.get('email'))
        if data.get('phone'):
            contact_items.append(data.get('phone'))
        if data.get('location'):
            contact_items.append(data.get('location'))
        
        contact_para.add_run(' • '.join(contact_items))
        
        # Social links
        if data.get('linkedin') or data.get('github') or data.get('portfolio'):
            social_para = doc.add_paragraph()
            
            if data.get('linkedin'):
                add_hyperlink(social_para, "LinkedIn", data.get('linkedin'))
                if data.get('github') or data.get('portfolio'):
                    social_para.add_run(" • ")
            
            if data.get('github'):
                add_hyperlink(social_para, "GitHub", data.get('github'))
                if data.get('portfolio'):
                    social_para.add_run(" • ")
            
            if data.get('portfolio'):
                add_hyperlink(social_para, "Portfolio", data.get('portfolio'))
        
        # Simple line separator
        separator = doc.add_paragraph()
        separator_run = separator.add_run("―" * 50)
        separator_run.font.color.rgb = RGBColor(0xC0, 0xC0, 0xC0)
        separator.space_after = Pt(12)
        
        return doc
    
    @staticmethod
    def _remove_table_borders(table):
        """Remove borders from table"""
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.first_child_found_in("w:tcBorders")
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)
                
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = tcBorders.find(qn(f'w:{border_name}'))
                    if border is not None:
                        tcBorders.remove(border)

def add_section_heading(doc, title, template_style="modern"):
    """Add a formatted section heading based on template style"""
    if template_style == "modern":
        heading = doc.add_heading(title, level=2)
        heading.runs[0].font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
        heading.runs[0].font.size = Pt(14)
        return heading
    elif template_style == "classic":
        heading = doc.add_heading(title, level=2)
        heading.runs[0].font.name = 'Times New Roman'
        heading.runs[0].font.size = Pt(12)
        heading.runs[0].font.bold = True
        return heading
    elif template_style == "creative":
        para = doc.add_paragraph()
        run = para.add_run(title.upper())
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
        return para
    else:  # minimal
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Calibri'
        para.space_before = Pt(12)
        para.space_after = Pt(6)
        return para

def create_safe_filename(name, document_type, template):
    """Create a safe filename from user input"""
    safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).rstrip()
    safe_name = safe_name.replace(' ', '_')
    if not safe_name:
        safe_name = "document"
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f"{document_type}_{safe_name}_{template}_{timestamp}.docx"

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/resume')
def resume_form():
    return render_template('resume_form.html')

@app.route('/cover-letter')
def cover_letter_form():
    return render_template('cover_letter_form.html')

@app.route('/generate-resume', methods=['POST'])
def generate_resume():
    try:
        data = request.json
        template_choice = data.get('template', 'modern')
        page_limit = data.get('page_limit', 1)
        
        # Create social links text for prompt
        social_links = []
        if data.get('linkedin'):
            social_links.append(f"LinkedIn: {data.get('linkedin')}")
        if data.get('github'):
            social_links.append(f"GitHub: {data.get('github')}")
        if data.get('portfolio'):
            social_links.append(f"Portfolio: {data.get('portfolio')}")
        
        social_links_text = " | ".join(social_links) if social_links else ""
        
        # Enhanced prompt with page limit consideration
        page_instruction = f"Keep content concise for a {page_limit}-page resume." if page_limit == 1 else f"You can use up to {page_limit} pages for detailed content."
        
        prompt = f"""
        Create a professional, ATS-friendly resume based on the following information. {page_instruction}
        
        Personal Information:
        - Name: {data.get('name', '')}
        - Email: {data.get('email', '')}
        - Phone: {data.get('phone', '')}
        - Location: {data.get('location', '')}
        - Professional Links: {social_links_text}
        
        Target Position: {data.get('job_title', '')}
        Years of Experience: {data.get('experience_years', '')}
        Industry: {data.get('industry', '')}
        
        Professional Summary: {data.get('professional_summary', '')}
        Professional Experience: {data.get('experience', '')}
        Skills: {data.get('skills', '')}
        Education: {data.get('education', '')}
        Certifications: {data.get('certifications', '')}
        Languages: {data.get('languages', '')}
        Career Objectives: {data.get('career_goals', '')}
        
        Requirements:
        1. Create well-structured sections appropriate for the content length
        2. Use professional language and action verbs
        3. Include quantifiable achievements
        4. Make it ATS-friendly
        5. Organize content logically
        6. {page_instruction}
        
        Structure with these sections:
        - Professional Summary (2-3 lines)
        - Core Skills (bullet points)
        - Professional Experience (reverse chronological)
        - Education
        - Additional sections if relevant (Certifications, Languages)
        """
        
        # Generate resume content
        resume_content = llm.generate_text(prompt, max_tokens=1500 if page_limit == 1 else 2500)
        
        # Create document with selected template
        doc = Document()
        
        # Apply template
        template_functions = {
            'modern': ResumeTemplates.create_modern_template,
            'classic': ResumeTemplates.create_classic_template,
            'creative': ResumeTemplates.create_creative_template,
            'minimal': ResumeTemplates.create_minimal_template
        }
        
        doc = template_functions[template_choice](doc, data, page_limit)
        
        # Add generated content with proper sections
        sections = resume_content.split('\n\n')
        current_template = template_choice
        
        for section in sections:
            if section.strip():
                # Check if it's a heading (contains common section words)
                section_keywords = ['SUMMARY', 'EXPERIENCE', 'SKILLS', 'EDUCATION', 'CERTIFICATIONS', 'LANGUAGES']
                is_heading = any(keyword in section.upper()[:30] for keyword in section_keywords)
                
                if is_heading:
                    # Extract heading and content
                    lines = section.split('\n')
                    heading = lines[0].strip()
                    content = '\n'.join(lines[1:]) if len(lines) > 1 else ""
                    
                    add_section_heading(doc, heading, current_template)
                    
                    if content:
                        para = doc.add_paragraph()
                        para.add_run(content)
                else:
                    para = doc.add_paragraph()
                    para.add_run(section)
        
        # Set page margins for better space utilization
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # Save document
        filename = create_safe_filename(data.get('name', 'user'), 'resume', template_choice)
        output_dir = Path('generated_documents')
        output_dir.mkdir(parents=True, exist_ok=True)
        filepath = output_dir / filename
        
        doc.save(str(filepath))
        
        if not filepath.exists():
            raise FileNotFoundError(f"Failed to create file: {filepath}")
        
        print(f"✅ Resume saved: {filepath.absolute()}")
        
        return jsonify({
            'success': True,
            'content': resume_content,
            'download_url': f'/download/{filename}',
            'filename': filename,
            'template': template_choice,
            'pages': page_limit
        })
        
    except Exception as e:
        print(f"❌ Error generating resume: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Failed to generate resume: {str(e)}'
        }), 500

@app.route('/generate-cover-letter', methods=['POST'])
def generate_cover_letter():
    try:
        data = request.json
        
        # Include social links in prompt
        social_links = []
        if data.get('linkedin'):
            social_links.append(f"LinkedIn: {data.get('linkedin')}")
        if data.get('github'):
            social_links.append(f"GitHub: {data.get('github')}")
        if data.get('portfolio'):
            social_links.append(f"Portfolio: {data.get('portfolio')}")
        
        social_links_text = " | ".join(social_links) if social_links else ""
        
        prompt = f"""
        Create a compelling, professional cover letter based on the following information:
        
        Applicant: {data.get('name', '')}
        Contact: {data.get('email', '')} | {data.get('phone', '')}
        Professional Links: {social_links_text}
        Target Company: {data.get('company', '')}
        Position: {data.get('position', '')}
        
        Job Requirements: {data.get('job_description', '')}
        Experience: {data.get('experience', '')}
        Skills: {data.get('skills', '')}
        Interest: {data.get('interest', '')}
        
        Create a personalized, engaging cover letter that highlights relevant experience.
        """
        
        # Generate content
        cover_letter_content = llm.generate_text(prompt, max_tokens=1200)
        
        # Create document
        doc = Document()
        
        # Header
        header = doc.add_paragraph()
        header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        header_run = header.add_run(f"{data.get('name', '')}\n")
        header_run.font.bold = True
        header_run.font.size = Pt(14)
        
        # Contact info
        contact_info = []
        if data.get('email'):
            contact_info.append(data.get('email'))
        if data.get('phone'):
            contact_info.append(data.get('phone'))
        
        if contact_info:
            header.add_run(" • ".join(contact_info) + "\n")
        
        # Social links with clickable hyperlinks
        if social_links:
            links_para = doc.add_paragraph()
            
            if data.get('linkedin'):
                add_hyperlink(links_para, "LinkedIn", data.get('linkedin'))
                if data.get('github') or data.get('portfolio'):
                    links_para.add_run(" | ")
            
            if data.get('github'):
                add_hyperlink(links_para, "GitHub", data.get('github'))
                if data.get('portfolio'):
                    links_para.add_run(" | ")
            
            if data.get('portfolio'):
                add_hyperlink(links_para, "Portfolio", data.get('portfolio'))
        
        # Date and recipient
        date_para = doc.add_paragraph()
        date_para.add_run(f"\n{datetime.now().strftime('%B %d, %Y')}\n\n")
        
        recipient = doc.add_paragraph()
        recipient.add_run(f"Hiring Manager\n{data.get('company', '')}\n\n")
        
        subject = doc.add_paragraph()
        subject_run = subject.add_run(f"Re: {data.get('position', '')} Position")
        subject_run.font.bold = True
        subject.add_run("\n\nDear Hiring Manager,\n")
        
        # Content
        content_para = doc.add_paragraph()
        content_para.add_run(cover_letter_content)
        
        # Closing
        closing = doc.add_paragraph()
        closing.add_run(f"\n\nBest regards,\n{data.get('name', '')}")
        
        # Save document
        filename = create_safe_filename(data.get('name', 'user'), 'cover_letter', 'standard')
        output_dir = Path('generated_documents')
        output_dir.mkdir(parents=True, exist_ok=True)
        filepath = output_dir / filename
        
        doc.save(str(filepath))
        
        if not filepath.exists():
            raise FileNotFoundError(f"Failed to create file: {filepath}")
        
        print(f"✅ Cover letter saved: {filepath.absolute()}")
        
        return jsonify({
            'success': True,
            'content': cover_letter_content,
            'download_url': f'/download/{filename}',
            'filename': filename
        })
        
    except Exception as e:
        print(f"❌ Error generating cover letter: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Failed to generate cover letter: {str(e)}'
        }), 500

@app.route('/download/<filename>')
def download_file(filename):
    try:
        safe_filename = os.path.basename(filename)
        filepath = Path('generated_documents') / safe_filename
        
        if not filepath.exists():
            return f"File not found: {safe_filename}", 404
        
        return send_file(str(filepath.absolute()), as_attachment=True)
        
    except Exception as e:
        return f"Error serving file: {str(e)}", 500

@app.route('/health')
def health_check():
    """Check if Ollama is running and model is available"""
    try:
        response = requests.get("http://localhost:11434/api/tags", timeout=5)
        if response.status_code == 200:
            models = response.json().get('models', [])
            return jsonify({
                'ollama_status': 'running',
                'available_models': [model['name'] for model in models]
            })
        else:
            return jsonify({'ollama_status': 'error', 'message': 'Ollama not responding'})
    except Exception as e:
        return jsonify({'ollama_status': 'offline', 'error': str(e)})

if __name__ == '__main__':
    print("🚀 Starting AI Resume & Cover Letter Generator...")
    print("📋 Make sure Ollama is running: ollama serve")
    print("🌐 Application will be available at: http://localhost:5000")
    print("❤️  Your data stays completely private on your machine!")
    app.run(debug=True, host='localhost', port=5000)
